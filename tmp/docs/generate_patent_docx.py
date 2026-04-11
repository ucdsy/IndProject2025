from __future__ import annotations

from pathlib import Path
from typing import Iterable
import subprocess
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path("/Users/xizhuxizhu/Desktop/IndProj04")
OUT = ROOT / "output" / "doc"
FIG_SVG_DIR = OUT / "专利正式附图版_20260410"
FIG_PNG_DIR = ROOT / "tmp" / "docs" / "figure_preview"


def set_cn_font(run, east_asia: str = "宋体", latin: str = "Times New Roman", size: int = 12, bold: bool = False) -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def apply_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.18)
    sec.right_margin = Cm(3.18)


def set_paragraph_format(paragraph, first_line_cm: float = 0.74, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=0, line_spacing=1.5) -> None:
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(first_line_cm) if first_line_cm else None
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = line_spacing
    paragraph.alignment = align


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, east_asia="黑体", size=16, bold=True)
    p.paragraph_format.space_after = Pt(12)


def add_center_name(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, east_asia="宋体", size=14, bold=False)
    p.paragraph_format.space_after = Pt(12)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_cm=0)
    r = p.add_run(text)
    set_cn_font(r, east_asia="黑体", size=14, bold=True)


def add_center_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_cn_font(r, east_asia="宋体", size=12)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p)
    r = p.add_run(text)
    set_cn_font(r, east_asia="宋体", size=12)


def add_list_item(doc: Document, text: str, bullet: str = "•") -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_cm=0, line_spacing=1.35)
    r = p.add_run(f"{bullet} {text}")
    set_cn_font(r, east_asia="宋体", size=12)


def add_number_item(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_cm=0, line_spacing=1.35)
    r = p.add_run(text)
    set_cn_font(r, east_asia="宋体", size=12)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)


def ensure_figure_png(svg_path: Path) -> Path:
    FIG_PNG_DIR.mkdir(parents=True, exist_ok=True)
    png_path = FIG_PNG_DIR / f"{svg_path.name}.png"
    if (
        png_path.exists()
        and png_path.stat().st_size > 0
        and png_path.stat().st_mtime >= svg_path.stat().st_mtime
    ):
        return png_path

    subprocess.run(
        [
            "/usr/bin/sips",
            "-s",
            "format",
            "png",
            str(svg_path),
            "--out",
            str(png_path),
        ],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    if not png_path.exists():
        raise FileNotFoundError(f"Failed to render preview for {svg_path}")
    return png_path


def add_figure_page(doc: Document, caption: str, svg_name: str, width_cm: float = 14.5) -> None:
    if caption:
        add_center_caption(doc, caption)
    png_path = ensure_figure_png(FIG_SVG_DIR / svg_name)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(png_path), width=Cm(width_cm))


def patent_figures() -> list[tuple[str, str]]:
    return [
        ("图1 总体流程示意图", "图1_总体流程示意图.svg"),
        ("图2 两层地址关系示意图", "图2_两层地址关系示意图.svg"),
        ("图3 异质共识与覆盖控制结构示意图", "图3_异质共识与覆盖控制结构示意图.svg"),
        ("图4 实例过滤与排序过程示意图", "图4_实例过滤与排序过程示意图.svg"),
        ("图5 结构化决策轨迹数据组织示意图", "图5_结构化决策轨迹数据组织示意图.svg"),
    ]


def parse_md_lines(path: Path) -> list[str]:
    return path.read_text(encoding="utf-8").splitlines()


def clean_markdown_text(text: str) -> str:
    cleaned = text.replace("`", "")
    while "](" in cleaned and "[" in cleaned:
        start = cleaned.find("[")
        mid = cleaned.find("](", start)
        end = cleaned.find(")", mid)
        if start == -1 or mid == -1 or end == -1:
            break
        label = cleaned[start + 1 : mid]
        cleaned = cleaned[:start] + label + cleaned[end + 1 :]
    return cleaned


def render_md_doc(doc: Document, lines: Iterable[str], title_mode: bool = True) -> None:
    first_nonempty = True
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        stripped = line.strip()

        if first_nonempty and title_mode:
            add_title(doc, stripped)
            first_nonempty = False
            continue

        if first_nonempty:
            first_nonempty = False

        if stripped.startswith("#"):
            add_heading(doc, stripped.lstrip("#").strip())
            continue

        if stripped.endswith("：") and len(stripped) <= 20:
            add_heading(doc, stripped[:-1])
            continue

        if stripped.startswith("- "):
            add_list_item(doc, clean_markdown_text(stripped[2:].strip()))
            continue

        # Remove markdown code ticks and file links for Word readability.
        add_body(doc, clean_markdown_text(stripped))


def build_spec_docx(src: Path, dst: Path) -> None:
    doc = Document()
    apply_normal_style(doc)
    render_md_doc(doc, parse_md_lines(src), title_mode=False)
    doc.save(dst)


def build_claims_docx(src: Path, dst: Path) -> None:
    doc = Document()
    apply_normal_style(doc)
    lines = parse_md_lines(src)
    first = True
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        stripped = line.strip()
        if first:
            add_title(doc, stripped)
            first = False
            continue
        add_body(doc, stripped)
    doc.save(dst)


def build_request_checklist_docx(src: Path, dst: Path) -> None:
    doc = Document()
    apply_normal_style(doc)
    lines = parse_md_lines(src)
    first = True
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        stripped = line.strip()
        if first:
            add_title(doc, stripped)
            first = False
            continue
        if stripped.endswith("：") and len(stripped) <= 20:
            add_heading(doc, stripped[:-1])
            continue
        if stripped.startswith("- "):
            add_list_item(doc, clean_markdown_text(stripped[2:].strip()))
            continue
        if stripped[:2].isdigit() and stripped[1] == ".":
            add_number_item(doc, clean_markdown_text(stripped))
            continue
        add_body(doc, clean_markdown_text(stripped))
    doc.save(dst)


def build_compilation_docx(dst: Path) -> None:
    doc = Document()
    apply_normal_style(doc)
    add_title(doc, "发明专利申请材料汇编")
    add_center_name(doc, "一种面向智能体语义寻址与执行发现的受约束多阶段路由决策方法、系统、设备及存储介质")
    add_body(doc, "本汇编版用于内部流转或提交专利代理人时集中查看，不替代分文件提交版本。")
    add_page_break(doc)

    sections = [
        ("说明书", OUT / "说明书初稿_20260406.md"),
        ("权利要求书", OUT / "权利要求书初稿_20260406.md"),
        ("说明书摘要", OUT / "说明书摘要初稿_20260406.md"),
        ("说明书附图清单与绘制说明", OUT / "说明书附图清单与绘制说明_20260406.md"),
        ("发明专利请求书著录项清单", OUT / "发明专利请求书著录项清单_20260406.md"),
    ]

    for idx, (section_title, path) in enumerate(sections):
        add_heading(doc, section_title)
        render_md_doc(doc, parse_md_lines(path), title_mode=False)
        if idx != len(sections) - 1:
            add_page_break(doc)

    doc.save(dst)


def build_figure_docx(dst: Path) -> None:
    doc = Document()
    apply_normal_style(doc)
    add_title(doc, "说明书附图")
    add_center_name(doc, "一种面向智能体语义寻址与执行发现的受约束多阶段路由决策方法、系统、设备及存储介质")

    figures = patent_figures()
    for idx, (caption, svg_name) in enumerate(figures):
        add_figure_page(doc, caption, svg_name)
        if idx != len(figures) - 1:
            add_page_break(doc)

    doc.save(dst)


def build_review_docx(dst: Path) -> None:
    doc = Document()
    apply_normal_style(doc)
    add_title(doc, "发明专利申请材料带图审阅版")
    add_center_name(doc, "一种面向智能体语义寻址与执行发现的受约束多阶段路由决策方法、系统、设备及存储介质")
    add_body(doc, "本版用于连续审阅。正式提交时仍按请求书、说明书、权利要求书、说明书摘要和说明书附图等文种分别整理。")
    add_page_break(doc)

    sections = [
        ("说明书", OUT / "说明书初稿_20260406.md"),
        ("权利要求书", OUT / "权利要求书初稿_20260406.md"),
        ("说明书摘要", OUT / "说明书摘要初稿_20260406.md"),
        ("说明书附图", None),
        ("发明专利请求书著录项清单", OUT / "发明专利请求书著录项清单_20260406.md"),
    ]

    for idx, (section_title, path) in enumerate(sections):
        add_heading(doc, section_title)
        if path is None:
            for figure_idx, (caption, svg_name) in enumerate(patent_figures()):
                add_figure_page(doc, caption, svg_name)
                if figure_idx != len(patent_figures()) - 1:
                    add_page_break(doc)
        else:
            render_md_doc(doc, parse_md_lines(path), title_mode=False)
        if idx != len(sections) - 1:
            add_page_break(doc)

    doc.save(dst)


def build_disclosure_docx(src: Path, dst: Path) -> None:
    doc = Document()
    apply_normal_style(doc)
    lines = parse_md_lines(src)

    first = True
    saw_figure_section = False
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        stripped = line.strip()

        if first:
            add_title(doc, stripped)
            first = False
            continue

        if stripped.startswith("申请："):
            add_body(doc, stripped)
            continue

        if stripped == "专利申请的名称：":
            add_heading(doc, stripped[:-1])
            continue

        if re.match(r"^\d+、", stripped):
            add_heading(doc, stripped)
            if stripped == "8、附图":
                saw_figure_section = True
            continue

        if stripped.startswith("- "):
            add_list_item(doc, clean_markdown_text(stripped[2:].strip()))
            continue

        if re.match(r"^\d+\.", stripped):
            add_number_item(doc, clean_markdown_text(stripped))
            continue

        add_body(doc, clean_markdown_text(stripped))

    if saw_figure_section:
        for caption, svg_name in patent_figures():
            add_center_caption(doc, caption)
            add_figure_page(doc, "", svg_name)

    doc.save(dst)


def main() -> None:
    targets = [
        (build_spec_docx, OUT / "说明书初稿_20260406.md", OUT / "说明书初稿_20260410.docx"),
        (build_claims_docx, OUT / "权利要求书初稿_20260406.md", OUT / "权利要求书初稿_20260410.docx"),
        (build_spec_docx, OUT / "说明书摘要初稿_20260406.md", OUT / "说明书摘要初稿_20260410.docx"),
        (build_spec_docx, OUT / "说明书附图清单与绘制说明_20260406.md", OUT / "说明书附图清单与绘制说明_20260410.docx"),
        (build_request_checklist_docx, OUT / "发明专利请求书著录项清单_20260406.md", OUT / "发明专利请求书著录项清单_20260410.docx"),
    ]

    for builder, src, dst in targets:
        builder(src, dst)

    build_compilation_docx(OUT / "发明专利申请材料汇编_20260410.docx")
    build_figure_docx(OUT / "说明书附图_20260410.docx")
    build_review_docx(OUT / "发明专利申请材料带图审阅版_20260410.docx")
    build_disclosure_docx(OUT / "专利技术交底书_20260410.md", OUT / "专利技术交底书_20260410.docx")


if __name__ == "__main__":
    main()
