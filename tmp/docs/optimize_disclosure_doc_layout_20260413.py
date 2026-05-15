from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path("/Users/xizhuxizhu/Desktop/IndProj04")
SRC = ROOT / "output/doc/专利技术交底书_20260410.docx"
OUT = ROOT / "output/doc/专利技术交底书_20260413_附图排版优化版.docx"


FIGURE_CAPTIONS = [
    "图1 总体流程示意图",
    "图2 两层地址关系示意图",
    "图3 异质共识与覆盖控制结构示意图",
    "图4 实例过滤与排序过程示意图",
    "图5 结构化决策轨迹数据组织示意图",
]

# Keep original embedded image content, but size portrait and landscape figures differently.
FIGURE_WIDTHS = [
    Inches(5.0),   # 图1 portrait
    Inches(6.2),   # 图2 landscape
    Inches(6.1),   # 图3 landscape
    Inches(6.2),   # 图4 landscape
    Inches(5.3),   # 图5 portrait-ish
]


def set_caption_format(paragraph, page_break_before=False):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.page_break_before = page_break_before
    fmt.keep_with_next = True
    fmt.space_before = Pt(12)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1.0
    for run in paragraph.runs:
        run.font.bold = True
        run.font.size = Pt(12)


def set_image_paragraph_format(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.keep_together = True
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(14)
    fmt.line_spacing = 1.0


def main():
    doc = Document(SRC)

    caption_indices = []
    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip() in FIGURE_CAPTIONS:
            caption_indices.append(idx)

    if len(caption_indices) != len(FIGURE_CAPTIONS):
        raise RuntimeError(f"Expected {len(FIGURE_CAPTIONS)} figure captions, found {len(caption_indices)}")

    # Format caption paragraphs and the immediately following image paragraphs.
    for pos, idx in enumerate(caption_indices):
        set_caption_format(doc.paragraphs[idx], page_break_before=True)
        if idx + 1 >= len(doc.paragraphs):
            raise RuntimeError(f"Caption paragraph {idx} has no following image paragraph")
        set_image_paragraph_format(doc.paragraphs[idx + 1])

    # Resize inline images in order of appearance.
    if len(doc.inline_shapes) != len(FIGURE_WIDTHS):
        raise RuntimeError(f"Expected {len(FIGURE_WIDTHS)} inline shapes, found {len(doc.inline_shapes)}")

    for shape, width in zip(doc.inline_shapes, FIGURE_WIDTHS):
        ratio = shape.height / shape.width
        shape.width = width
        shape.height = int(width * ratio)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
