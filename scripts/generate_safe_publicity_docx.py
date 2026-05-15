from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


INPUT_MD = Path("output/doc/项目成果宣传稿_安全公开版_20260414.md")
OUTPUT_DOCX = Path("output/doc/项目成果宣传稿_安全公开版_20260414.docx")


def apply_font(run, size: int, east_asia: str, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def add_body_paragraph(doc: Document, text: str, indent: bool = True) -> None:
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.84)
    run = p.add_run(text)
    apply_font(run, 12, "宋体")


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        apply_font(run, 16, "黑体", bold=True)
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        apply_font(run, 14, "黑体", bold=True)


def build_doc(lines: list[str]) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    in_postscript = False
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line == "---":
            in_postscript = True
            continue
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), level=2)
            continue

        if in_postscript and (line[:2].isdigit() or (line[0].isdigit() and line[1:2] == ".")):
            add_body_paragraph(doc, line, indent=False)
            continue

        if line.endswith("：") or line.endswith(":"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(line)
            apply_font(run, 12, "黑体", bold=True)
            continue

        add_body_paragraph(doc, line, indent=True)

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.25

    return doc


def main() -> None:
    text = INPUT_MD.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = build_doc(lines)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
