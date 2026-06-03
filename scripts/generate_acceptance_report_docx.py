from copy import deepcopy
import argparse
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


TEMPLATE = Path(
    "/Users/xizhuxizhu/Library/Containers/com.tencent.xinWeChat/Data/Documents/"
    "xwechat_files/wxid_tt5mguhxte5g22_44ce/temp/drag/自立科研项目验收报告模板.docx"
)
INPUT_MD = Path("output/doc/自立科研项目验收报告_填写版_20260426.md")
OUTPUT_DOCX = Path("output/doc/自立科研项目验收报告_填写版_20260426.docx")


def set_font(run, size=12, east_asia="宋体", bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def set_paragraph_text(paragraph, text, size=12, east_asia="宋体", bold=False, align=None):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_font(run, size=size, east_asia=east_asia, bold=bold)
    paragraph.paragraph_format.line_spacing = 1.5
    if align is not None:
        paragraph.alignment = align


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    new_para._element = new_p
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(paragraph, rows, cols):
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Cm(16))
    paragraph._p.addnext(table._tbl)
    return table


def insert_image_after(paragraph, image_path, caption):
    caption_para = insert_paragraph_after(paragraph, caption)
    set_paragraph_text(
        caption_para,
        caption,
        size=10.5,
        east_asia="宋体",
        bold=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    caption_para.paragraph_format.line_spacing = 1.0

    image_para = insert_paragraph_after(paragraph)
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.paragraph_format.line_spacing = 1.0
    run = image_para.add_run()
    run.add_picture(str(image_path), width=Cm(15.2))
    return image_para


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def parse_markdown(path):
    sections = {}
    cover = {}
    current = None
    buffer = []
    lines = path.read_text(encoding="utf-8").splitlines()
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("# "):
            continue
        if current is None and (
            stripped.startswith("项目")
            or stripped.startswith("承研")
            or stripped.startswith("起止")
            or stripped.startswith("填报")
        ):
            if "：" in stripped:
                k, v = stripped.split("：", 1)
                cover[k] = v
            continue
        if stripped.startswith("## "):
            if current:
                sections[current] = buffer
            current = stripped[3:].strip()
            buffer = []
            continue
        if current is not None:
            buffer.append(line.rstrip())
    if current:
        sections[current] = buffer
    return cover, sections


def add_body_after(anchor, lines):
    cursor = anchor
    pending_table = []
    in_table = False

    def flush_table():
        nonlocal cursor, pending_table
        if not pending_table:
            return
        rows = list(reversed([row for row in pending_table if row]))
        if len(rows) >= 2:
            table = insert_table_after(anchor, len(rows), len(rows[0]))
            try:
                table.style = "Table Grid"
            except KeyError:
                pass
            set_table_borders(table)
            for r_idx, row in enumerate(rows):
                for c_idx, value in enumerate(row):
                    cell = table.cell(r_idx, c_idx)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    run = p.add_run(value)
                    set_font(run, 10.5, "宋体", bold=(r_idx == 0))
                    p.paragraph_format.line_spacing = 1.2
            cursor = anchor
        pending_table = []

    for raw in reversed(lines):
        line = raw.strip()
        if not line:
            continue
        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(re.fullmatch(r"-+", c.replace(" ", "")) for c in cells):
                continue
            pending_table.append(cells)
            in_table = True
            continue
        if in_table:
            flush_table()
            in_table = False
        image_match = re.match(r"^!\[(.+?)\]\((.+?)\)$", line)
        if image_match:
            caption = image_match.group(1).strip()
            image_path = Path(image_match.group(2).strip())
            if not image_path.is_absolute():
                image_path = Path.cwd() / image_path
            if image_path.exists():
                insert_image_after(anchor, image_path, caption)
            else:
                para = insert_paragraph_after(anchor, f"{caption}（图片文件待补：{image_path}）")
                set_paragraph_text(para, para.text, size=12, east_asia="宋体")
            continue
        if line.startswith("### "):
            para = insert_paragraph_after(anchor, line[4:].strip())
            set_paragraph_text(para, line[4:].strip(), size=12, east_asia="黑体", bold=True)
            cursor = para
        elif re.match(r"^\\d+\\.\\s", line):
            para = insert_paragraph_after(anchor, line)
            set_paragraph_text(para, line, size=12, east_asia="宋体")
            para.paragraph_format.first_line_indent = Cm(0)
            cursor = para
        else:
            para = insert_paragraph_after(anchor, line)
            set_paragraph_text(para, line, size=12, east_asia="宋体")
            para.paragraph_format.first_line_indent = Cm(0.84)
            cursor = para
    if in_table:
        flush_table()


def iter_all_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def fill_cover(doc, cover):
    replacements = {
        "项目编号：": f"项目编号：{cover.get('项目编号', '')}",
        "项目名称：": f"项目名称：{cover.get('项目名称', '')}",
        "项目类型：": f"项目类型：{cover.get('项目类型', '')}",
        "承研处/所：": f"承研处/所：{cover.get('承研处/所', '')}",
        "项目负责人：": f"项目负责人：{cover.get('项目负责人', '')}",
        "起止时间：": f"起止时间：{cover.get('起止时间', '')}",
        "项目经费：": f"项目经费：{cover.get('项目经费', '')}",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            set_paragraph_text(paragraph, replacements[text], size=14, east_asia="宋体")
        elif text.startswith("填报日期"):
            set_paragraph_text(paragraph, f"填报日期  {cover.get('填报日期', '')}", size=12, east_asia="宋体", align=WD_ALIGN_PARAGRAPH.CENTER)


def remove_instruction_paragraphs(doc):
    template_only_headings = {
        "2.1 目标完成情况",
        "2.2 主要工作内容",
        "2.3 考核指标完成情况",
        "2.4 项目亮点",
        "5.1实施进度、质量、人员投入等方面",
        "5.2产学研用合作情况",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if (
            text.startswith("（自立科研项目任务书")
            or text.startswith("（详细阐述项目目标")
            or text.startswith("（项目成果被")
            or text in template_only_headings
        ):
            clear_paragraph(paragraph)


def fill_sections(doc, sections):
    section_map = {
        "一、项目概况": "一、项目概况",
        "二、项目完成情况": "二、项目完成情况",
        "三、成果转化应用情况": "三、成果转化应用情况",
        "四、经费使用情况": "四、经费使用情况",
        "五、组织管理情况": "五、组织管理情况",
        "六、存在问题及建议": "六、存在问题及建议",
        "七、有关证明材料": "七、有关证明材料",
    }
    anchors = {p.text.strip(): p for p in doc.paragraphs if p.text.strip() in section_map}
    for title, source_title in section_map.items():
        anchor = anchors.get(title)
        if anchor and source_title in sections:
            add_body_after(anchor, sections[source_title])


def style_existing_headings(doc):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if re.match(r"^[一二三四五六七八]、", text):
            set_paragraph_text(paragraph, text, size=14, east_asia="黑体", bold=True)
        elif re.match(r"^\\d\\.\\d", text):
            set_paragraph_text(paragraph, text, size=12, east_asia="黑体", bold=True)


def build_fallback_template(cover):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.6)

    def add_line(text="", size=12, east_asia="宋体", bold=False, align=None):
        p = doc.add_paragraph()
        set_paragraph_text(p, text, size=size, east_asia=east_asia, bold=bold, align=align)
        return p

    add_line(f"项目编号：{cover.get('项目编号', '')}", size=12)
    for _ in range(5):
        add_line("")
    add_line("自立科研项目验收报告", size=18, east_asia="黑体", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(4):
        add_line("")
    add_line(f"项目名称：{cover.get('项目名称', '')}", size=14)
    add_line(f"项目类型：{cover.get('项目类型', '')}", size=14)
    add_line(f"承研处/所：{cover.get('承研处/所', '')}", size=14)
    add_line(f"项目负责人：{cover.get('项目负责人', '')}", size=14)
    add_line(f"起止时间：{cover.get('起止时间', '')}", size=14)
    add_line(f"项目经费：{cover.get('项目经费', '')}", size=14)
    for _ in range(3):
        add_line("")
    add_line(f"填报日期  {cover.get('填报日期', '')}", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_line("中国互联网络信息中心制", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    for title in (
        "一、项目概况",
        "二、项目完成情况",
        "三、成果转化应用情况",
        "四、经费使用情况",
        "五、组织管理情况",
        "六、存在问题及建议",
        "七、有关证明材料",
        "八、审批意见",
    ):
        add_line(title, size=14, east_asia="黑体", bold=True)

    table = doc.add_table(rows=4, cols=2)
    set_table_borders(table)
    rows = (
        ("项目负责人意见", "项目负责人签字：\n\n          年   月   日"),
        ("承研处所负责人意见", "承研处所负责人签字：\n\n          年   月   日"),
        ("分管领导意见", "分管领导签字：\n\n          年   月   日"),
        ("组织实施单位意见", "（加盖中心公章）\n\n          年   月   日"),
    )
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(value)
            set_font(run, 12, "宋体")
    return doc


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input-md", type=Path, default=INPUT_MD)
    parser.add_argument("--output-docx", type=Path, default=OUTPUT_DOCX)
    args = parser.parse_args()

    cover, sections = parse_markdown(args.input_md)
    if TEMPLATE.exists():
        doc = Document(TEMPLATE)
        fill_cover(doc, cover)
        remove_instruction_paragraphs(doc)
    else:
        doc = build_fallback_template(cover)
    fill_sections(doc, sections)
    style_existing_headings(doc)

    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output_docx)
    print(args.output_docx)


if __name__ == "__main__":
    main()
