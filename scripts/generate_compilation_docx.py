import re
import shutil
import tempfile
from zipfile import ZIP_DEFLATED, ZipFile
from xml.etree import ElementTree as ET
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


INPUT_MD = Path("output/doc/项目成果汇编_20260416.md")
OUTPUT_DOCX = Path("output/doc/项目成果汇编_20260416.docx")

FOOTNOTES = {
    "[fn:air-canada]": (
        "Moffatt v. Air Canada, 2024 BCCRT 149, Civil Resolution Tribunal "
        "(British Columbia), CanLII: "
        "https://www.canlii.org/en/bc/bccrt/doc/2024/2024bccrt149/2024bccrt149.html。"
    )
}
FOOTNOTE_IDS = {marker: i for i, marker in enumerate(FOOTNOTES, start=1)}
USED_FOOTNOTE_IDS: set[int] = set()


def apply_font(
    run,
    size: int,
    east_asia: str,
    bold: bool = False,
    latin: str | None = None,
) -> None:
    latin_font = latin or east_asia
    run.font.name = latin_font
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:ascii"), latin_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin_font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:cs"), latin_font)


def add_text_with_footnotes(
    paragraph,
    text: str,
    size: int,
    east_asia: str,
    bold: bool = False,
    latin: str | None = None,
) -> None:
    pattern = "(" + "|".join(re.escape(marker) for marker in FOOTNOTES) + ")"
    for part in re.split(pattern, text):
        if not part:
            continue
        if part in FOOTNOTE_IDS:
            footnote_id = FOOTNOTE_IDS[part]
            USED_FOOTNOTE_IDS.add(footnote_id)
            run = paragraph.add_run()
            apply_font(run, size, east_asia, bold=bold, latin=latin)
            ref = OxmlElement("w:footnoteReference")
            ref.set(qn("w:id"), str(footnote_id))
            run._r.append(ref)
            continue
        if not bold and east_asia == "仿宋":
            add_math_text(paragraph, part, size=size)
        else:
            run = paragraph.add_run(part)
            apply_font(run, size, east_asia, bold=bold, latin=latin)


def add_math_text(paragraph, text: str, size: int = 16) -> None:
    for token in re.split(r"([A-Za-zΑ-Ωα-ωΠπΔδΦφΚκΡρΣσΖζ]+(?:_[A-Za-z0-9]+)+)", text):
        if not token:
            continue
        if "_" not in token:
            run = paragraph.add_run(token)
            apply_font(run, size, "仿宋")
            continue

        parts = token.split("_")
        base_run = paragraph.add_run(parts[0])
        apply_font(base_run, size, "仿宋")
        for sub in parts[1:]:
            sub_run = paragraph.add_run(sub)
            apply_font(sub_run, size, "仿宋")
            sub_run.font.subscript = True


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    style = doc.styles["Normal"]
    style.font.name = "仿宋"
    style.font.size = Pt(16)
    style._element.rPr.rFonts.set(qn("w:ascii"), "仿宋")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "仿宋")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    style._element.rPr.rFonts.set(qn("w:cs"), "仿宋")


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    apply_font(run, 22, "方正小标宋简体", latin="Times New Roman")


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level == 2 else 6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(1.13)
    run = p.add_run(text)
    if level == 2:
        apply_font(run, 16, "黑体")
    elif level == 3:
        apply_font(run, 16, "楷体")
    else:
        apply_font(run, 16, "仿宋")


def add_body_paragraph(doc: Document, text: str, indent: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.13)
    p.paragraph_format.space_after = Pt(0)
    add_text_with_footnotes(p, text, 16, "仿宋")


def add_formula_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)

    label, _, expression = text.partition("：")
    label_run = p.add_run(label + "：")
    apply_font(label_run, 16, "仿宋")
    add_math_text(p, expression.rstrip("。"), size=16)
    if text.endswith("。"):
        end_run = p.add_run("。")
        apply_font(end_run, 16, "仿宋")


def add_tag_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(1.13)
    add_text_with_footnotes(p, text, 16, "仿宋")


def add_image(doc: Document, caption: str, path_text: str) -> None:
    image_path = Path(path_text.strip())
    if not image_path.is_absolute():
        image_path = Path.cwd() / image_path
    if not image_path.exists():
        raise FileNotFoundError(f"Image not found: {image_path}")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(15.2))

    if caption:
        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_p.paragraph_format.space_after = Pt(4)
        caption_run = caption_p.add_run(caption)
        apply_font(caption_run, 16, "仿宋")


def add_table_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    apply_font(run, 16, "仿宋")


def parse_table_row(line: str) -> list[str]:
    if not line.strip().startswith("|"):
        return []
    return [cell.strip().replace("`", "") for cell in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    cells = parse_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return

    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.autofit = True
    table.allow_autofit = True

    for row_index, row in enumerate(rows):
        for col_index in range(cols):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            text = row[col_index] if col_index < len(row) else ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 or col_index > 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.1
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            apply_font(run, 10, "仿宋", bold=row_index == 0)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def build_doc(lines: list[str]) -> Document:
    doc = Document()
    configure_document(doc)

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            add_title(doc, line[2:].strip())
            i += 1
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("#### "):
            add_heading(doc, line[5:].strip(), level=4)
            i += 1
            continue
        image_match = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            add_image(doc, image_match.group(1).strip(), image_match.group(2).strip())
            i += 1
            continue
        if (
            line.startswith("|")
            and i + 1 < len(lines)
            and is_table_separator(lines[i + 1].strip())
        ):
            table_rows = [parse_table_row(line)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_line = lines[i].strip()
                if not is_table_separator(table_line):
                    table_rows.append(parse_table_row(table_line))
                i += 1
            add_markdown_table(doc, table_rows)
            continue
        if re.match(r"^表\d+[\s　]", line):
            add_table_caption(doc, line)
            i += 1
            continue
        if re.match(r"^公式\d+：", line):
            add_formula_paragraph(doc, line)
            i += 1
            continue
        if line.startswith("项目编号：") or line.startswith("研究处所：") or line.startswith("项目负责人：") or line.startswith("项目参与人："):
            add_tag_paragraph(doc, line)
            i += 1
            continue
        add_body_paragraph(doc, line, indent=True)
        i += 1

    return doc


def build_footnotes_xml(used_ids: set[int]) -> bytes:
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ET.register_namespace("w", ns)
    root = ET.Element(f"{{{ns}}}footnotes")

    sep = ET.SubElement(root, f"{{{ns}}}footnote", {f"{{{ns}}}type": "separator", f"{{{ns}}}id": "-1"})
    p = ET.SubElement(sep, f"{{{ns}}}p")
    r = ET.SubElement(p, f"{{{ns}}}r")
    ET.SubElement(r, f"{{{ns}}}separator")

    cont = ET.SubElement(root, f"{{{ns}}}footnote", {f"{{{ns}}}type": "continuationSeparator", f"{{{ns}}}id": "0"})
    p = ET.SubElement(cont, f"{{{ns}}}p")
    r = ET.SubElement(p, f"{{{ns}}}r")
    ET.SubElement(r, f"{{{ns}}}continuationSeparator")

    for marker, footnote_id in FOOTNOTE_IDS.items():
        if footnote_id not in used_ids:
            continue
        footnote = ET.SubElement(root, f"{{{ns}}}footnote", {f"{{{ns}}}id": str(footnote_id)})
        p = ET.SubElement(footnote, f"{{{ns}}}p")
        p_pr = ET.SubElement(p, f"{{{ns}}}pPr")
        ET.SubElement(p_pr, f"{{{ns}}}pStyle", {f"{{{ns}}}val": "FootnoteText"})
        ref_r = ET.SubElement(p, f"{{{ns}}}r")
        r_pr = ET.SubElement(ref_r, f"{{{ns}}}rPr")
        ET.SubElement(r_pr, f"{{{ns}}}rStyle", {f"{{{ns}}}val": "FootnoteReference"})
        ET.SubElement(ref_r, f"{{{ns}}}footnoteRef")
        text_r = ET.SubElement(p, f"{{{ns}}}r")
        text_r_pr = ET.SubElement(text_r, f"{{{ns}}}rPr")
        fonts = ET.SubElement(text_r_pr, f"{{{ns}}}rFonts")
        fonts.set(f"{{{ns}}}ascii", "仿宋")
        fonts.set(f"{{{ns}}}hAnsi", "仿宋")
        fonts.set(f"{{{ns}}}eastAsia", "仿宋")
        ET.SubElement(text_r_pr, f"{{{ns}}}sz", {f"{{{ns}}}val": "21"})
        text = ET.SubElement(text_r, f"{{{ns}}}t")
        text.text = FOOTNOTES[marker]

    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def inject_footnotes(docx_path: Path, used_ids: set[int]) -> None:
    if not used_ids:
        return

    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    ET.register_namespace("", rel_ns)
    ET.register_namespace("", ct_ns)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = Path(tmp.name)

    with ZipFile(docx_path, "r") as zin, ZipFile(tmp_path, "w", ZIP_DEFLATED) as zout:
        names = set(zin.namelist())
        rels_xml = zin.read("word/_rels/document.xml.rels")
        rels_root = ET.fromstring(rels_xml)
        has_footnotes_rel = any(
            rel.attrib.get("Type") == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
            for rel in rels_root
        )
        if not has_footnotes_rel:
            rids = []
            for rel in rels_root:
                rid = rel.attrib.get("Id", "")
                if rid.startswith("rId") and rid[3:].isdigit():
                    rids.append(int(rid[3:]))
            next_rid = max(rids or [0]) + 1
            ET.SubElement(
                rels_root,
                f"{{{rel_ns}}}Relationship",
                {
                    "Id": f"rId{next_rid}",
                    "Type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
                    "Target": "footnotes.xml",
                },
            )

        ct_xml = zin.read("[Content_Types].xml")
        ct_root = ET.fromstring(ct_xml)
        has_footnotes_ct = any(
            child.attrib.get("PartName") == "/word/footnotes.xml"
            for child in ct_root
        )
        if not has_footnotes_ct:
            ET.SubElement(
                ct_root,
                f"{{{ct_ns}}}Override",
                {
                    "PartName": "/word/footnotes.xml",
                    "ContentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
                },
            )

        for item in zin.infolist():
            if item.filename in {"word/footnotes.xml", "word/_rels/document.xml.rels", "[Content_Types].xml"}:
                continue
            zout.writestr(item, zin.read(item.filename))
        zout.writestr("word/footnotes.xml", build_footnotes_xml(used_ids))
        zout.writestr("word/_rels/document.xml.rels", ET.tostring(rels_root, encoding="utf-8", xml_declaration=True))
        zout.writestr("[Content_Types].xml", ET.tostring(ct_root, encoding="utf-8", xml_declaration=True))

    shutil.move(tmp_path, docx_path)


def main() -> None:
    USED_FOOTNOTE_IDS.clear()
    text = INPUT_MD.read_text(encoding="utf-8")
    doc = build_doc(text.splitlines())
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    inject_footnotes(OUTPUT_DOCX, USED_FOOTNOTE_IDS)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
