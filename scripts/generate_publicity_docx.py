from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image, ImageDraw, ImageFont


TITLE = "系列研究成果（拟）：面向互联网基础资源的大模型多智能体协作与可信认知标识技术研究"
INTRO = (
    "本文围绕“面向互联网基础资源的大模型多智能体协作与可信认知标识技术研究”项目，"
    "聚焦互联网基础资源场景中的智能体命名与语义路由这一典型方向，针对能力表达方式不统一、"
    "自然语言请求难以稳定映射到可调用能力、复杂约束场景下路由结果缺乏可解释支撑等问题，"
    "开展以大模型语义理解为基础、以多智能体协同复核为增强、以结构化过程记录为支撑的高可信语义路由研究，"
    "构建能力命名空间、结构化语义路由与受控协作复核框架，完成原型系统设计和系列实验验证工作。"
)

SECTIONS = [
    (
        "一、互联网基础资源场景下的智能体命名与地址表达研究",
        [
            "借鉴互联网基础资源分层命名与稳定标识思路，对互联网基础资源相关智能体能力进行系统梳理，构建面向语义路由的能力命名与地址表达方式。相关命名体系既可用于语义路由任务的统一标注，也可用于候选组织、原型验证和成果展示。",
            "在地址表达方面，区分了面向语义路由的能力地址和面向执行发现的实例地址。能力地址用于表达请求应交由哪一类能力处理，实例地址用于表达由哪一个具体智能体实例执行。通过将能力寻址与执行落点分层组织，为智能体注册、发现和调用提供了统一地址基础。",
        ],
    ),
    (
        "二、互联网基础资源场景下的高可信语义路由原型框架",
        [
            "围绕自然语言请求到执行服务的处理过程，构建了由候选召回、规则基线路由、结构化语义路由、受控协作复核和执行落点解析组成的高可信语义路由原型框架。系统首先在固定命名空间内组织候选集合，再在候选集合内部完成主能力与相关能力判断，并在必要时触发协作复核，最终完成从能力地址到执行实例的服务落点。该框架强调在受限候选边界内开展决策，而不是进行开放式生成，从而使召回错误、候选内裁决错误和执行落点错误可以被清晰拆解。",
            "在该框架中，系统依托大模型对复杂请求进行语义理解、竞争候选辨析和结构化判断生成，在复杂边界样本上结合多智能体协同复核机制增强路由判断的稳定性和准确性，并对候选组织、裁决、升级、改判和执行落点等关键环节进行结构化记录，形成可解释、可复盘、可核验的处理轨迹，为可信认知标识技术在语义路由场景中的应用提供支撑。",
            "围绕复杂任务表达，系统在能力命名空间内围绕节点描述、别名、层级信息和元数据标签组织候选集合，在此基础上综合主任务命中、上下文匹配、层级粒度和节点类型等信息完成直接路由。进一步地，围绕主能力判断、相关能力恢复、竞争候选说明和不确定性摘要形成结构化决策，并与规则分数进行受控融合，使系统不仅能够给出路由结果，还能够说明为何选择某一能力、为何未选择竞争候选，以及哪些因素可能影响最终改判。",
            "在此基础上，围绕复杂样本进一步建立受控协作复核机制。协作复核并非自由讨论，而是基于职责分工组织不同视角的复核过程，并通过显式放行策略控制最终是否改判。与此同时，对候选组织、裁决、升级、改判和执行落点等关键环节进行结构化记录，形成可回放、可解释、可复核的过程轨迹，为可信认知标识技术在语义路由场景下的落地应用提供了基础。",
        ],
    ),
    (
        "三、原型系统与实验验证",
        [
            "在上述研究基础上，形成了面向智能体命名与语义路由的原型系统，能够对自然语言请求、候选组织、路由判断、协作复核和执行落点等关键环节进行全过程展示。原型系统既可用于典型场景演示，也可用于路由结果解释、过程复盘和后续能力扩展。",
            "围绕原型框架，研究团队开展了多轮、多批次实验验证。系列实验结果显示，在直接路由基础上引入多智能体协同复核后，复杂场景下的路由准确率进一步提升，说明以大模型语义理解为基础、以多智能体协作为增强的处理方式，能够有效增强智能体命名与语义路由任务中的判断稳定性和结果可靠性。相关结果同时表明，高可信语义路由能力提升的关键不在于简单增加更多复核轮次，而在于结构化语义证据、职责化复核和显式改判授权的协同作用；结构化过程记录则进一步增强了处理过程的可解释性、可复盘性和可核验性。",
            "本研究围绕项目中智能体命名与语义路由这一互联网基础资源典型场景，构建了高可信语义路由原型框架，形成了可展示、可验证、可复盘的原型系统；通过系列实验验证了大模型多智能体协同处理对路由准确率提升的积极作用；通过结构化过程记录体现了可信认知标识技术在结果解释、过程复盘和可核验支撑方面的应用价值；同步形成了技术报告、论文材料、专利材料、演示脚本和答辩展示材料等成果，为推动互联网基础资源向智能体命名、寻址和可信调用方向延伸提供了有益探索。",
        ],
    ),
]


def pick_font():
    candidates = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    return None


def build_figure(path: Path) -> None:
    width, height = 1800, 520
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font_path = pick_font()
    font = ImageFont.truetype(font_path, 34) if font_path else ImageFont.load_default()
    font_small = ImageFont.truetype(font_path, 28) if font_path else ImageFont.load_default()

    boxes = [
        ("自然语言请求", (70, 120, 270, 220)),
        ("能力命名空间", (310, 120, 540, 220)),
        ("候选召回", (580, 120, 760, 220)),
        ("直接路由", (800, 120, 980, 220)),
        ("结构化语义路由", (1020, 120, 1280, 220)),
        ("受控协作复核", (1320, 120, 1580, 220)),
        ("执行落点解析", (1620, 120, 1790, 220)),
    ]
    trace_box = ("过程留痕与可信认知标识支撑", (420, 330, 1380, 420))

    for text, (x1, y1, x2, y2) in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill="#EEF4FF", outline="#4A6FA5", width=3)
        bbox = draw.multiline_textbbox((0, 0), text, font=font, align="center")
        tw = bbox[2] - bbox[0]
        th = bbox[3] - bbox[1]
        draw.multiline_text(
            (x1 + (x2 - x1 - tw) / 2, y1 + (y2 - y1 - th) / 2 - 2),
            text,
            font=font,
            fill="#1F1F1F",
            align="center",
        )

    text, (x1, y1, x2, y2) = trace_box
    draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill="#EFFAF2", outline="#4B8E62", width=3)
    bbox = draw.multiline_textbbox((0, 0), text, font=font, align="center")
    tw = bbox[2] - bbox[0]
    th = bbox[3] - bbox[1]
    draw.multiline_text(
        (x1 + (x2 - x1 - tw) / 2, y1 + (y2 - y1 - th) / 2 - 2),
        text,
        font=font,
        fill="#1F1F1F",
        align="center",
    )

    for i in range(len(boxes) - 1):
        _, (x1, y1, x2, y2) = boxes[i]
        _, (nx1, ny1, _, ny2) = boxes[i + 1]
        y = (y1 + y2) // 2
        start = (x2 + 5, y)
        end = (nx1 - 10, y)
        draw.line([start, end], fill="#6B7280", width=5)
        draw.polygon([(end[0], end[1]), (end[0] - 18, end[1] - 10), (end[0] - 18, end[1] + 10)], fill="#6B7280")

    for _, (x1, y1, x2, y2) in [boxes[2], boxes[3], boxes[4], boxes[5], boxes[6]]:
        sx = (x1 + x2) // 2
        sy = y2 + 6
        ey = trace_box[1][1] - 8
        steps = 12
        for j in range(steps):
            y_start = sy + j * ((ey - sy) / steps)
            y_end = y_start + 8
            draw.line([(sx, y_start), (sx, y_end)], fill="#A0A0A0", width=3)

    title = "大模型多智能体协作与可信认知标识支撑下的高可信语义路由原型框架"
    bbox = draw.textbbox((0, 0), title, font=font_small)
    tw = bbox[2] - bbox[0]
    draw.text(((width - tw) / 2, 36), title, font=font_small, fill="#374151")
    img.save(path)


def apply_font(run, size: int, east_asia: str, bold: bool = False) -> None:
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def main():
    out_dir = Path("output/doc")
    out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / "项目成果宣传稿_公众号版.docx"
    fig_path = out_dir / "项目成果宣传稿_公众号版_架构图.png"
    build_figure(fig_path)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    apply_font(run, 16, "黑体", bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.84)
    run = p.add_run(INTRO)
    apply_font(run, 12, "宋体")

    for idx, (heading, paragraphs) in enumerate(SECTIONS):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(heading)
        apply_font(run, 14, "黑体", bold=True)

        for para_idx, paragraph in enumerate(paragraphs):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.84)
            run = p.add_run(paragraph)
            apply_font(run, 12, "宋体")

            if idx == 1 and para_idx == 1:
                pic = doc.add_paragraph()
                pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic.add_run().add_picture(str(fig_path), width=Inches(6.3))

                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap.add_run("图1 大模型多智能体协作与可信认知标识支撑下的高可信语义路由原型框架")
                apply_font(run, 10, "宋体")

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.25

    doc.save(docx_path)
    print(docx_path)


if __name__ == "__main__":
    main()
